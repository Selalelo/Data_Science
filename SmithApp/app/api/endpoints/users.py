from fastapi import APIRouter, Depends, HTTPException, status
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from core.database import get_db
from schemas.user import UserCreate, UserUpdate, UserResponse
from services.user_service import UserService
from api.deps import get_current_user
from typing import List
from docx import Document
from docx.shared import Inches
from io import BytesIO

router = APIRouter()


@router.get("/", response_model=List[UserResponse])
def get_all_users(
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Get all users
    Requires authentication
    """
    users = UserService.get_all_users(db)
    return users


@router.get("/{user_id}", response_model=UserResponse)
def get_user(
    user_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Get user by ID
    Requires authentication
    """
    user = UserService.get_user_by_id(db, user_id)
    
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    return user


@router.post("/", response_model=UserResponse, status_code=status.HTTP_201_CREATED)
def create_user(
    user_data: UserCreate,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Create a new user
    Requires authentication
    """
    try:
        user = UserService.create_user(db, user_data)
        return user
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Failed to create user: {str(e)}"
        )


@router.put("/{user_id}", response_model=UserResponse)
def update_user(
    user_id: int,
    user_data: UserUpdate,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Update user details (FirstName and LastName only)
    Requires authentication
    """
    user = UserService.update_user(db, user_id, user_data)
    
    if not user:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    return user


@router.delete("/{user_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_user(
    user_id: int,
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Delete user
    ONLY allows User ID 1 (Admin) to perform this action
    """
    # 1. Check if the logged-in user is User ID 1
    # Note: Depending on your current_user structure, it might be current_user["user_id"] or current_user.user_id
    if current_user.get("user_id") != 1:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="Permission denied: Only the primary administrator can delete users."
        )

    # 2. Prevent user from deleting themselves (optional but recommended)
    if user_id == current_user.get("user_id"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="You cannot delete your own administrator account."
        )

    # 3. Proceed with deletion
    success = UserService.delete_user(db, user_id)
    
    if not success:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="User not found"
        )
    
    return None



@router.get("/export/word", response_class=StreamingResponse)
def export_users_to_word(
    db: Session = Depends(get_db),
    current_user: dict = Depends(get_current_user)
):
    """
    Export users list to MS Word document
    Requires authentication
    """
    users = UserService.get_all_users(db)
    
    # Create Word document
    doc = Document()
    doc.add_heading('User Management Report', 0)
    
    # Add table
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Name'
    header_cells[1].text = 'Date of Birth'
    header_cells[2].text = 'Province'
    header_cells[3].text = 'Gender'
    header_cells[4].text = 'Facilitator'
    
    # Add data rows
    for user in users:
        row_cells = table.add_row().cells
        row_cells[0].text = user['full_name']
        row_cells[1].text = str(user['date_of_birth'])
        row_cells[2].text = user['province']
        row_cells[3].text = user['gender']
        row_cells[4].text = 'Yes' if user['facilitator'] else 'No'
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=users_report.docx"}
    )